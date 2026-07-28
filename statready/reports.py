from __future__ import annotations

from io import BytesIO
import json
import zipfile
from typing import Any

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter

from .literature import references_for_method
from .models import AnalysisResult, AuditEntry


def audit_frame(entries: list[AuditEntry]) -> pd.DataFrame:
    return pd.DataFrame([entry.to_dict() for entry in entries]) if entries else pd.DataFrame(
        columns=["created_at", "action", "variable", "details", "justification", "before_n", "after_n"]
    )


DISPLAY_COLUMN_NAMES = {
    "std_error": "SE",
    "statistic": "Statistic",
    "p_value": "p-value",
    "ci_lower": "CI lower",
    "ci_upper": "CI upper",
    "adjusted_r_squared": "Adjusted R-squared",
    "r_squared": "R-squared",
    "f_statistic": "F statistic",
    "f_p_value": "F p-value",
    "selected_penalty_alpha": "Selected ridge alpha",
    "cross_validation_folds": "CV folds",
    "cross_validated_rmse": "CV RMSE",
    "cross_validated_r_squared": "CV R-squared",
    "training_r_squared": "Training R-squared",
    "predictor_columns": "Predictor columns",
    "ridge_coefficient_per_1_sd_predictor": "Ridge coefficient per 1 SD",
    "standardized_coefficient": "Standardised coefficient",
    "absolute_standardized_coefficient": "Absolute coefficient",
    "importance_rank": "Rank",
    "ols_standardized_coefficient": "OLS standardised coefficient",
    "ridge_standardized_coefficient": "Ridge standardised coefficient",
    "direction_agrees": "Direction agrees",
    "data_values_changed": "Data values changed",
    "degrees_of_freedom": "Degrees of freedom",
    "chi_square": "Chi-square",
    "chi_square_p": "Chi-square p-value",
    "standardized_loading": "Standardised loading",
    "standardized_estimate": "Standardised estimate",
    "composite_reliability": "Composite reliability",
    "average_variance_extracted": "Average variance extracted",
    "greenhouse_geisser_epsilon": "Greenhouse-Geisser epsilon",
    "greenhouse_geisser_corrected_p": "GG corrected p-value",
    "intraclass_correlation": "Intraclass correlation",
    "hausman_chi_square": "Hausman chi-square",
    "hausman_p": "Hausman p-value",
}


def _format_value(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def _set_cell_text(cell, text: str, font_size: float = 8.5, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(font_size)


def _set_table_row_layout(row, repeat_header: bool = False) -> None:
    """Keep table rows intact and repeat header rows across page breaks."""
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        properties.append(OxmlElement("w:tblHeader"))


def _add_dataframe(document: Document, frame: pd.DataFrame, max_rows: int = 100) -> None:
    if frame.empty:
        document.add_paragraph("No results available.")
        return

    display = frame.head(max_rows).copy()
    display = display.rename(columns=DISPLAY_COLUMN_NAMES)

    # Wide single-record outputs are much clearer as vertical key-value tables.
    if len(display) <= 2 and len(display.columns) > 5:
        for record_number, (_, row) in enumerate(display.iterrows(), start=1):
            if len(display) > 1:
                document.add_paragraph(f"Record {record_number}").runs[0].bold = True
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            _set_table_row_layout(table.rows[0], repeat_header=True)
            _set_cell_text(table.rows[0].cells[0], "Measure", 8.5, True)
            _set_cell_text(table.rows[0].cells[1], "Value", 8.5, True)
            for column, value in row.items():
                new_row = table.add_row()
                _set_table_row_layout(new_row)
                cells = new_row.cells
                _set_cell_text(cells[0], str(column), 8.5, True)
                _set_cell_text(cells[1], _format_value(value), 8.5)
        return

    # Split very wide multi-record outputs into readable metric blocks. The first
    # identifier column is repeated in every block so that rows remain traceable.
    if len(display) > 2 and len(display.columns) > 10:
        identifier = display.columns[0]
        measure_columns = list(display.columns[1:])
        chunk_size = 6
        for start in range(0, len(measure_columns), chunk_size):
            chunk = measure_columns[start:start + chunk_size]
            label = ", ".join(str(column) for column in chunk)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Measures: {label}")
            run.bold = True
            run.font.size = Pt(8.5)
            _add_dataframe(document, display[[identifier] + chunk], max_rows=max_rows)
        if len(frame) > max_rows:
            document.add_paragraph(
                f"Only the first {max_rows} of {len(frame)} rows are shown in this document. "
                "The Excel export contains the full table."
            )
        return

    table = document.add_table(rows=1, cols=len(display.columns))
    table.style = "Table Grid"
    _set_table_row_layout(table.rows[0], repeat_header=True)
    font_size = 7.2 if len(display.columns) >= 7 else 8.2 if len(display.columns) >= 5 else 9
    for idx, column in enumerate(display.columns):
        _set_cell_text(table.rows[0].cells[idx], str(column), font_size, True)
    for _, row in display.iterrows():
        new_row = table.add_row()
        _set_table_row_layout(new_row)
        cells = new_row.cells
        for idx, value in enumerate(row):
            _set_cell_text(cells[idx], _format_value(value), font_size)
    if len(frame) > max_rows:
        document.add_paragraph(f"Only the first {max_rows} of {len(frame)} rows are shown in this document. The Excel export contains the full table.")


def build_docx_report(
    result: AnalysisResult,
    study: dict[str, Any],
    analysis_plan: pd.DataFrame,
    audit_entries: list[AuditEntry],
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading(study.get("title") or "StatReady Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Phase 2.6 guided statistical analysis, validated AI planning, diagnostics and reproducibility report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Study specification", level=1)
    for label, key in [
        ("Objective", "objective"),
        ("Hypothesis", "hypothesis"),
        ("Significance level", "alpha"),
        ("Selected method", "method"),
    ]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(str(study.get(key, "")))

    document.add_heading("Analysis plan", level=1)
    _add_dataframe(document, analysis_plan)

    descriptive_tables = {
        name: table for name, table in result.tables.items()
        if name.startswith("Descriptive ")
    }
    inferential_tables = {
        name: table for name, table in result.tables.items()
        if not name.startswith("Descriptive ")
    }

    if descriptive_tables:
        document.add_heading("Descriptive statistics", level=1)
        document.add_paragraph(
            str(result.metadata.get("descriptive_summary") or
                "Descriptive statistics summarise the variables used in the analysis before inferential results are interpreted.")
        )
        descriptive_display_names = {
            "Descriptive sample overview": "Analysis sample overview",
            "Descriptive statistics - Numeric variables": "Numeric variables",
            "Descriptive statistics - Categorical summary": "Categorical summary",
            "Descriptive statistics - Frequencies": "Frequencies",
            "Descriptive statistics - By group": "Descriptive statistics by group",
            "Descriptive profile overview": "Profile sample overview",
            "Descriptive profile - Numeric variables": "Profile numeric variables",
            "Descriptive profile - Categorical summary": "Profile categorical summary",
            "Descriptive profile - Frequencies": "Profile frequencies",
        }
        for name, table in descriptive_tables.items():
            clean_name = descriptive_display_names.get(
                name, name.replace("Descriptive statistics - ", "").replace("Descriptive ", "")
            )
            document.add_heading(clean_name, level=2)
            if name == "Descriptive statistics - Numeric variables":
                centre_columns = [
                    column for column in [
                        "variable", "valid_n", "missing_n", "missing_percent",
                        "mean", "std_dev", "median",
                    ] if column in table.columns
                ]
                shape_columns = [
                    column for column in [
                        "variable", "minimum", "q1", "q3", "maximum",
                        "skewness", "kurtosis",
                    ] if column in table.columns
                ]
                document.add_paragraph("Sample, centre and dispersion").runs[0].bold = True
                _add_dataframe(document, table[centre_columns])
                document.add_paragraph("Range and distribution shape").runs[0].bold = True
                _add_dataframe(document, table[shape_columns])
            else:
                _add_dataframe(document, table)

    document.add_heading("Inferential results summary" if inferential_tables else "Results summary", level=1)
    document.add_paragraph(result.summary)
    if result.warnings:
        document.add_heading("Warnings", level=2)
        for warning in result.warnings:
            document.add_paragraph(warning, style="List Bullet")

    if result.figures:
        is_network = "network" in result.method.lower()
        document.add_heading("Network analysis figures" if is_network else "Path and measurement diagram", level=1)
        document.add_paragraph(
            "The figures present network structure, communities, centrality, degree distribution and adjacency patterns. Interpret them with the construction rule, diagnostics and stability tables."
            if is_network else
            "The diagram presents the prespecified model and the standardised estimates obtained from the fitted analysis. It should be interpreted together with the coefficient tables, diagnostics and fit indices."
        )
        for figure_name, figure_bytes in result.figures.items():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(BytesIO(figure_bytes), width=Inches(6.8))
            caption = document.add_paragraph(f"Figure: {figure_name}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption.runs:
                caption.runs[0].italic = True

    compact_row_limits = {
        "Latent variable scores": 25,
        "Cluster sizes": 30,
        "Cluster residual influence screening": 30,
        "Cross-loadings": 40,
        "Factor scores": 25,
    }
    for name, table in inferential_tables.items():
        if name == "Paper-ready abstract":
            document.add_page_break()
        document.add_heading(name, level=2)
        if name.startswith("Paper-ready") and "text" in table.columns:
            for _, row in table.iterrows():
                if row.get("section"):
                    paragraph = document.add_paragraph()
                    paragraph.add_run(str(row.get("section"))).bold = True
                document.add_paragraph(str(row.get("text", "")))
        else:
            _add_dataframe(document, table, max_rows=compact_row_limits.get(name, 100))

    document.add_heading("Diagnostics and assumptions", level=1)
    if result.diagnostics.empty:
        document.add_paragraph("No method-specific diagnostics were generated.")
    else:
        diagnostic_core = [column for column in ["diagnostic", "test", "statistic", "p_value", "status"] if column in result.diagnostics.columns]
        _add_dataframe(document, result.diagnostics[diagnostic_core])
        narrative_columns = [column for column in ["diagnostic", "interpretation", "recommended_response"] if column in result.diagnostics.columns]
        if len(narrative_columns) >= 2:
            document.add_heading("Diagnostic interpretation and response", level=2)
            _add_dataframe(document, result.diagnostics[narrative_columns])
        concern_count = int(result.diagnostics["status"].isin(["Minor concern", "Material concern"]).sum()) if "status" in result.diagnostics else 0
        document.add_paragraph(f"The diagnostic review identified {concern_count} item(s) requiring interpretation or sensitivity analysis. A failed diagnostic does not authorise alteration of observations to obtain statistical significance.")

    if result.metadata.get("diagnostic_response"):
        document.add_heading("Diagnostic response and alternative model", level=2)
        document.add_paragraph(str(result.metadata["diagnostic_response"]))

    combined_audit = audit_entries + result.treatment_log
    document.add_heading("Data-treatment and analysis audit trail", level=1)
    _add_dataframe(document, audit_frame(combined_audit))

    references = references_for_method(result.method, result.diagnostics)
    document.add_heading("Supporting methodological literature", level=1)
    for _, row in references.iterrows():
        text = row["citation"]
        if row.get("doi"):
            text += f" DOI: {row['doi']}"
        document.add_paragraph(text, style="List Bullet")

    document.add_page_break()
    document.add_heading("Reproducibility code", level=1)
    code_paragraph = document.add_paragraph()
    code_run = code_paragraph.add_run(result.reproducible_code or "No code generated.")
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(5.5)
    code_paragraph.paragraph_format.space_after = Pt(0)
    code_paragraph.paragraph_format.line_spacing = 0.75

    document.add_paragraph(
        "Integrity statement: The report distinguishes the original data, documented data preparation and analysis-specific sensitivity procedures. No treatment should be applied solely to obtain a preferred p-value or conclusion."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_excel_report(
    original_df: pd.DataFrame,
    analysis_df: pd.DataFrame,
    result: AnalysisResult,
    analysis_plan: pd.DataFrame,
    audit_entries: list[AuditEntry],
) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        original_df.to_excel(writer, sheet_name="Original_Data", index=False)
        analysis_df.to_excel(writer, sheet_name="Analysis_Data", index=False)
        analysis_plan.to_excel(writer, sheet_name="Analysis_Plan", index=False)
        if not result.diagnostics.empty:
            result.diagnostics.to_excel(writer, sheet_name="Diagnostics", index=False)
        audit_frame(audit_entries + result.treatment_log).to_excel(writer, sheet_name="Audit_Trail", index=False)
        references_for_method(result.method, result.diagnostics).to_excel(writer, sheet_name="Method_References", index=False)
        result_sheet_names = {
            "Descriptive sample overview": "Descriptive_Overview",
            "Descriptive statistics - Numeric variables": "Descriptive_Numeric",
            "Descriptive statistics - Categorical summary": "Descriptive_Categories",
            "Descriptive statistics - Frequencies": "Descriptive_Frequencies",
            "Descriptive statistics - By group": "Descriptive_By_Group",
            "Descriptive profile overview": "Profile_Overview",
            "Descriptive profile - Numeric variables": "Profile_Numeric",
            "Descriptive profile - Categorical summary": "Profile_Categories",
            "Descriptive profile - Frequencies": "Profile_Frequencies",
            "Selected coefficient table": "Selected_Coefficients",
            "Conventional coefficients": "OLS_Coefficients",
            "HC3 robust coefficients": "HC3_Coefficients",
            "VIF": "VIF",
            "Multicollinearity action summary": "Collinearity_Action",
            "Predictor correlation matrix": "Predictor_Correlations",
            "Ridge sensitivity coefficients": "Ridge_Coefficients",
            "Ridge sensitivity model fit": "Ridge_Fit",
            "OLS-Ridge coefficient comparison": "OLS_Ridge_Compare",
            "Factor loadings": "EFA_Loadings",
            "Factor variance explained": "EFA_Variance",
            "Parallel analysis": "EFA_Parallel",
            "KMO by item": "EFA_KMO",
            "CFA fit indices": "CFA_Fit",
            "CFA standardised loadings": "CFA_Loadings",
            "Construct reliability and validity": "CFA_Reliability",
            "SEM fit indices": "SEM_Fit",
            "Structural path estimates": "SEM_Paths",
            "SEM standardised loadings": "SEM_Loadings",
            "Repeated-measures ANOVA": "RM_ANOVA",
            "Holm-adjusted pairwise comparisons": "RM_Pairwise",
            "Fixed effects": "Mixed_Fixed",
            "Mixed-model fit": "Mixed_Fit",
            "Panel model decision": "Panel_Decision",
            "Selected panel coefficients": "Panel_Selected",
            "Parallel indirect effects": "Parallel_Indirect",
            "Conditional indirect effects": "Conditional_Indirect",
            "Model fit": "Model_Fit",
        }
        for idx, (name, table) in enumerate(result.tables.items(), start=1):
            safe_name = result_sheet_names.get(name)
            if safe_name is None:
                safe_name = "".join(ch for ch in name if ch not in r"[]:*?/\\")[:25]
            sheet = f"R{idx}_{safe_name}"[:31]
            table.to_excel(writer, sheet_name=sheet, index=False)

        for figure_index, (figure_name, figure_bytes) in enumerate(result.figures.items(), start=1):
            safe_base = "".join(ch for ch in figure_name if ch.isalnum() or ch == "_")[:22] or "Figure"
            sheet_name = f"Fig{figure_index}_{safe_base}"[:31]
            pd.DataFrame({
                "Diagram": [figure_name],
                "Interpretation": ["Standardised estimates. Interpret together with the path/loading tables and model-fit diagnostics."],
            }).to_excel(writer, sheet_name=sheet_name, index=False)
            worksheet = writer.sheets[sheet_name]
            image = XLImage(BytesIO(figure_bytes))
            max_width = 1180
            if image.width > max_width:
                scale = max_width / image.width
                image.width = int(image.width * scale)
                image.height = int(image.height * scale)
            worksheet.add_image(image, "A4")
            worksheet.column_dimensions["A"].width = 34
            worksheet.column_dimensions["B"].width = 72
            worksheet.row_dimensions[2].height = 36

        header_fill = PatternFill("solid", fgColor="17365D")
        header_font = Font(color="FFFFFF", bold=True)
        concern_fill = PatternFill("solid", fgColor="F4CCCC")
        for worksheet in writer.book.worksheets:
            worksheet.freeze_panes = "A2"
            worksheet.auto_filter.ref = worksheet.dimensions
            worksheet.sheet_view.showGridLines = False
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            worksheet.row_dimensions[1].height = 28

            for column_index, column_cells in enumerate(worksheet.iter_cols(), start=1):
                header = str(column_cells[0].value or "").lower()
                max_length = len(str(column_cells[0].value or ""))
                for cell in column_cells[1:]:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    value = cell.value
                    if value is not None:
                        max_length = max(max_length, len(str(value)))
                    if isinstance(value, float):
                        if "percent" in header:
                            cell.number_format = "0.00"
                        elif "p_value" in header or "p-value" in header:
                            cell.number_format = "0.0000"
                        else:
                            cell.number_format = "0.000"
                    if isinstance(value, int):
                        cell.number_format = "0"
                width = min(max(max_length + 2, 11), 38)
                if header in {"variable", "term", "category", "diagnostic", "test", "status"}:
                    width = min(max(width, 16), 28)
                if header in {"details", "justification", "interpretation", "recommended_response", "interpretive_decision", "specification", "citation", "supports"}:
                    width = 38
                worksheet.column_dimensions[get_column_letter(column_index)].width = width

            headers = {str(cell.value): cell.column for cell in worksheet[1] if cell.value is not None}
            status_column = headers.get("status")
            if status_column:
                for row in range(2, worksheet.max_row + 1):
                    if str(worksheet.cell(row, status_column).value) == "Material concern":
                        worksheet.cell(row, status_column).fill = concern_fill
                        worksheet.cell(row, status_column).font = Font(color="9C0006", bold=True)
    return buffer.getvalue()


def build_reproducibility_package(
    original_df: pd.DataFrame,
    analysis_df: pd.DataFrame,
    result: AnalysisResult,
    study: dict[str, Any],
    analysis_plan: pd.DataFrame,
    audit_entries: list[AuditEntry],
) -> bytes:
    docx = build_docx_report(result, study, analysis_plan, audit_entries)
    xlsx = build_excel_report(original_df, analysis_df, result, analysis_plan, audit_entries)
    audit_csv = audit_frame(audit_entries + result.treatment_log).to_csv(index=False).encode("utf-8")
    references_csv = references_for_method(result.method, result.diagnostics).to_csv(index=False).encode("utf-8")
    metadata = {
        "study": study,
        "analysis_method": result.method,
        "summary": result.summary,
        "metadata": result.metadata,
        "warnings": result.warnings,
    }

    package = BytesIO()
    with zipfile.ZipFile(package, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("StatReady_Report.docx", docx)
        archive.writestr("StatReady_Results.xlsx", xlsx)
        archive.writestr("Original_Data.csv", original_df.to_csv(index=False))
        archive.writestr("Analysis_Data.csv", analysis_df.to_csv(index=False))
        archive.writestr("Treatment_Audit.csv", audit_csv)
        archive.writestr("Methodological_References.csv", references_csv)
        archive.writestr("Reproduce_Analysis.py", result.reproducible_code)
        for figure_name, figure_bytes in result.figures.items():
            safe_name = "".join(ch if ch.isalnum() else "_" for ch in figure_name).strip("_")
            archive.writestr(f"Figures/{safe_name}.png", figure_bytes)
        if result.metadata.get("interactive_network_html"):
            archive.writestr("Figures/Interactive_Network.html", str(result.metadata["interactive_network_html"]))
        if (result.metadata.get("diagram_settings") or {}).get("custom_positions"):
            archive.writestr("Figures/Path_Diagram_Arrangement.json", json.dumps(result.metadata["diagram_settings"]["custom_positions"], indent=2))
        archive.writestr("Analysis_Metadata.json", json.dumps(metadata, indent=2, default=str))
    return package.getvalue()
